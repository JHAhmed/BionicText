import re
import docx
import customtkinter
from docx.shared import Pt as fontSize
from math import floor as roundDown
from tkinter.filedialog import askopenfilename, asksaveasfilename

customtkinter.set_appearance_mode("dark")
customtkinter.set_default_color_theme("dark-blue")

root = customtkinter.CTk()
root.geometry("380x420")
root.title("Bionic Text")

selectedFileName = ""
selectedFile = ""
newDoc = ""

def UpdateLabel () :
    global label
    label.configure(text = "You've selected: " + selectedFileName)

def UpdateSavedLabel () :
    global savedLabel
    savedLabel.configure(text = "File Saved Successfully!")

def selectFontSize_callback(choice):
    global label
    # label.configure(text = "Font Size: " + str(choice))
    print(str(choice))

def selectFont_callback(choice):
    global label
    # label.configure(text = "Font Size: " + str(choice))
    print(str(choice))

def ChooseFile () :
    global selectedFile
    global selectedFileName
    selectedFile = askopenfilename()
    selectedFileList = selectedFile.split("/")
    selectedFileName = selectedFileList[-1]
    UpdateLabel()

def SaveFile () :
    filename = asksaveasfilename()
    return filename

def ConvertToBionic (para) :
    global newDoc
    totalText = []
    tempList = []
    tempList.append(re.split(r'(\s+)', para))

    for everyList in tempList :
        for word in everyList :
            totalText.append(word)

    doc_para = newDoc.add_paragraph()

    for word in totalText:
        charToBold = roundDown(len(word)/2)
        tempSplit = list(word)

        for i in range(charToBold):
            doc_para.add_run(tempSplit[i]).bold = True

        for i in range(charToBold, len(tempSplit)) :
            doc_para.add_run(tempSplit[i])

    tempList.clear()  
    tempSplit.clear()

def StartProcess () :
    global fontVar
    global fontSizeVar

    global newDoc
    doc = docx.Document(selectedFile)
    newDoc = docx.Document()

    totalParas = []

    font = newDoc.styles['Normal'].font
    font.name = str(fontVar.get())
    font.size = fontSize(int(str(fontSizeVar.get())))

    for para in doc.paragraphs:
        totalParas.append(para.text);

    for para in totalParas:
        if (len(para) > 0) :
            ConvertToBionic(para)
        
    newDoc.save(str(SaveFile()) + ".docx")
    UpdateSavedLabel()

frame = customtkinter.CTkFrame(master = root)
frame.pack(pady = 20, padx = 60, fill = "both", expand = True)

chooseButton = customtkinter.CTkButton(master = frame, text = "Choose File", command = ChooseFile)
chooseButton.pack(pady = 12, padx = 10)

label = customtkinter.CTkLabel(master = frame, text = "You've selected: ")
label.pack(pady = 12, padx = 10)

fontSizeVar = customtkinter.StringVar(value = "12")
fontVar = customtkinter.StringVar(value = "Calibri")

fontSizeValues = ["8", "10", "12", "14", "16", "18", "20", "24", "28"]
selectFontSize = customtkinter.CTkOptionMenu(master = frame, values = fontSizeValues,
                                             command = selectFontSize_callback, variable = fontSizeVar)
selectFontSize.pack(padx = 20, pady = 10)

fontValues = ["Calibri", "Times New Roman", "Arial"]
selectFont = customtkinter.CTkOptionMenu(master = frame, values = fontValues,
                                         command = selectFont_callback, variable = fontVar)
selectFont.pack(padx = 20, pady = 10)

saveButton = customtkinter.CTkButton(master = frame, text = "Convert To Bionic", command = StartProcess)
saveButton.pack(pady = 12, padx = 10)

savedLabel = customtkinter.CTkLabel(master = frame, text = "...")
savedLabel.pack(pady = 12, padx = 10)

def exitProgram():
    global root
    root.destroy()
    quit()

quitButton = customtkinter.CTkButton(master = frame, text = "Quit", command = exitProgram)
quitButton.pack(pady = 12, padx = 10)

root.mainloop()